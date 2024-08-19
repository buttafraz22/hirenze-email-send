from jinja2 import Template
import smtplib
import os
import shutil
import re
from docx2pdf import convert
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from docx import Document


def clear_dir(pdf_dir):
    if os.path.exists(pdf_dir):
        for filename in os.listdir(pdf_dir):
            file_path = os.path.join(pdf_dir, filename)
            if os.path.isfile(file_path):
                os.remove(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
            

def docx_replace_regex(doc_obj, regex , replace):
    """A function to replace patterns of regex in the 
    
    """
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)

def populate_html(template_literal:str, **kwargs)->str:
    template = Template(template_literal)
    return template.render(kwargs)


def create_docx(context: dict)->str:
    """A function to create certificate of a single person in a pdf and docx
    form.

    Args:
        context (dict): A dictionary of keys mapping to each person. These keys are: 
            - name(str): The name of the recipient of the certificate.
            - event(str): The name of the event of the certificate.
            - ambassador(str): The ambassador whose event it was.
    
    """
    try:
        os.makedirs('Output/pdf')
    except:
        ...

    doc = Document('doc-template.docx')

    
    reg = re.compile(r""+'{Name Surname}')
    replace = r""+context['name']   # John Doe
    docx_replace_regex(doc, reg, replace)

    reg = re.compile(r""+'{EVENT NAME}')
    replace = r""+context['event'] # Serverless Surge
    docx_replace_regex(doc, reg, replace)

    reg = re.compile(r""+'{AMBASSADOR NAME}')
    replace = r""+context['ambassador'] # Hammad Hassan
    docx_replace_regex(doc, reg, replace)
    
    outstream = f'Output/pdf/Cert-{context["name"]}.docx'
    doc.save(outstream)

    convert(outstream)
    os.remove(outstream) # Delete the docx
    return f'Output/pdf/Cert-{context["name"]}.pdf'

def send_email(
        req_body:dict
    )->None:
    """A function to create the email and send it to users.

    Args:
        req_body (dict): A set of request parameters to be sent in the email. Format:
            - req_body= {
                'emails': ['list of emails'],
                'names': ['list of names, index between email and name shared'],
                'ambassador': 'ambassador-name',
                'ambassador_email': 'sender-email',
                'ambassador_password': 'your-app-password',
                'event': 'event name',
                'subject': 'Certificate of participation',
                'html_template' : an HTML template for name
            }    
    """

    server = 'smtp.gmail.com'
    smtp_port = 587
    server = smtplib.SMTP(server, smtp_port)
    server.starttls()
    server.login(req_body['ambassador_email'],req_body['ambassador_password'])

    
    for email, name in zip(req_body['emails'], req_body['names']):
        
        msg = MIMEMultipart('mixed')
        msg['From'] = req_body['ambassador_email']
        msg['To'] = email
        msg['Subject'] = req_body['subject']

        html_context = req_body.copy()
        del html_context['html_template']; del html_context['ambassador_password']
        del html_context['ambassador_email']; del html_context['subject']
        del html_context['names']; del html_context['emails']
        html_context['name'] = name

        

        # Populate HTML content with recipient's name using Jinja2
        populated_html = populate_html(req_body['html_template'], **html_context)


        # Attach HTML to the email
        msg.attach(MIMEText(populated_html, 'html'))

        # Attach the certificate
        certificate_context = {
            'name': name,
            'ambassador': req_body['ambassador'],
            'event': req_body['event'] 
            }
        pdf_path = create_docx(context=certificate_context)
        with open(pdf_path, 'rb') as file:
            attachment = MIMEApplication(file.read(), _subtype='pdf')
            attachment.add_header('Content-Disposition', 'attachment', filename=f"Cert-{name}.pdf")
            msg.attach(attachment)

        # Send email
        server.sendmail(req_body['ambassador_email'], email, msg.as_string())

    # Quit SMTP server
    server.quit()


if __name__=='__main__':
    from template import template_str
    req_body= {
        'emails': [''],  # You are solely responsible for keeping the length same
        'names': [''],   # of names and emails, and shared indexing. I would add csv loading soon.
        'ambassador': '',
        'ambassador_email': '', # Supply email
        'ambassador_password': '',  # Supply password
        'event': 'Title of event',
        'subject': 'Certificate of participation',
        'html_template' : template_str  
    }

    send_email(req_body)
    clear_dir('Output/pdf')

